from datetime import datetime
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportGenerator:
    def __init__(self):
        self.report_date = datetime.now().strftime('%Y-%m-%d')
        self.output_dir = os.path.expanduser('~/Desktop/OSINT_Reports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_text_report(self, findings, output_path):
        target = findings.get('target', 'Unknown')
        
        if DOCX_AVAILABLE:
            return self._generate_docx(findings)
        else:
            return self._generate_txt(findings)
    
    def _generate_docx(self, findings):
        doc = Document()
        
        target = findings.get('target', 'Unknown')
        summary = findings.get('summary', {})
        risk = summary.get('risk_assessment', {})
        risk_scores = summary.get('risk_scores', {})
        red_flags = summary.get('red_flags', [])
        
        # Title
        title = doc.add_heading('BLACK WIDOW GLOBAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('CONFIDENTIAL INTELLIGENCE DOSSIER', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Target Info
        doc.add_heading(f'Subject: {target}', level=1)
        doc.add_paragraph(f"Report Date: {self.report_date}")
        doc.add_paragraph(f"Classification: CONFIDENTIAL - FOR CLIENT USE ONLY")
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        
        level = risk.get('level', 'UNKNOWN')
        score = risk.get('score', 0)
        signal = risk.get('invest_signal', 'UNKNOWN')
        rec = risk.get('recommendation', '')
        
        exec_para = doc.add_paragraph()
        exec_para.add_run(f"RISK ASSESSMENT: {level}\n").bold = True
        exec_para.add_run(f"RISK SCORE: {score:.0f}/100\n")
        exec_para.add_run(f"INVESTMENT SIGNAL: {signal}\n").bold = True
        exec_para.add_run(f"\nRECOMMENDATION: {rec}")
        
        doc.add_paragraph()
        
        # Risk Matrix Table
        doc.add_heading('RISK MATRIX', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Category'
        header_cells[1].text = 'Score'
        header_cells[2].text = 'Assessment'
        
        for cat, score_val in risk_scores.items():
            row_cells = table.add_row().cells
            row_cells[0].text = cat
            row_cells[1].text = str(score_val)
            if score_val >= 60:
                row_cells[2].text = 'HIGH RISK'
            elif score_val >= 40:
                row_cells[2].text = 'ELEVATED'
            elif score_val >= 25:
                row_cells[2].text = 'MODERATE'
            else:
                row_cells[2].text = 'LOW'
        
        doc.add_paragraph()
        
        # Critical Findings
        if red_flags:
            doc.add_heading('CRITICAL FINDINGS', level=1)
            
            for flag in red_flags:
                doc.add_heading(f"[{flag.get('severity', 'MEDIUM')}] {flag.get('category', 'General')}", level=2)
                
                finding_para = doc.add_paragraph()
                finding_para.add_run("Finding: ").bold = True
                finding_para.add_run(flag.get('finding', ''))
                
                impact_para = doc.add_paragraph()
                impact_para.add_run("Business Impact: ").bold = True
                impact_para.add_run(flag.get('so_what', ''))
                
                action_para = doc.add_paragraph()
                action_para.add_run("Recommended Action: ").bold = True
                action_para.add_run(flag.get('action', ''))
                
                doc.add_paragraph()
        
        # Intelligence Analysis
        doc.add_heading('INTELLIGENCE ANALYSIS', level=1)
        
        news_data = findings.get('data_sources', {}).get('news_search', {})
        if isinstance(news_data, dict):
            articles = news_data.get('articles', [])
            adverse = news_data.get('adverse_media', [])
            
            if adverse:
                doc.add_heading('Adverse Media Coverage', level=2)
                for article in adverse[:10]:
                    para = doc.add_paragraph()
                    para.add_run(f"• {article.get('title', 'Unknown')}\n").bold = True
                    para.add_run(f"  Source: {article.get('domain', 'Unknown')} | Date: {article.get('date', 'Unknown')}\n")
                    if article.get('url'):
                        para.add_run(f"  Link: {article.get('url', '')}")
                    doc.add_paragraph()
        
        # Data Sources
        doc.add_heading('DATA SOURCES CONSULTED', level=1)
        
        sources = [
            "SEC EDGAR (Securities and Exchange Commission filings)",
            "Federal Court Records (CourtListener/RECAP)",
            "FEC (Federal Election Commission political donations)",
            "GDELT Global News Database",
            "OpenSanctions (Global sanctions and watchlists)",
            "LinkedIn Company Intelligence",
            "State Corporate Registries"
        ]
        
        for source in sources:
            doc.add_paragraph(f"• {source}")
        
        doc.add_paragraph()
        
        # Disclaimer
        doc.add_heading('DISCLAIMER', level=1)
        disclaimer = doc.add_paragraph()
        disclaimer.add_run(
            "This report is provided for informational purposes only and does not constitute legal, financial, or investment advice. "
            "Black Widow Global makes no warranties regarding the completeness or accuracy of this information. "
            "All information is derived from publicly available sources. Recipients should independently verify all findings "
            "before making business decisions. This report is confidential and intended solely for the use of the client."
        )
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("BLACK WIDOW GLOBAL | Corporate Intelligence & Investigative Due Diligence").bold = True
        footer.add_run(f"\nReport Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Save
        safe_name = target.replace(' ', '_').replace(',', '').replace('.', '')[:30]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Intel_Report_{safe_name}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _generate_txt(self, findings):
        target = findings.get('target', 'Unknown')
        summary = findings.get('summary', {})
        risk = summary.get('risk_assessment', {})
        
        lines = [
            "=" * 70,
            "BLACK WIDOW GLOBAL - CONFIDENTIAL INTELLIGENCE DOSSIER",
            "=" * 70,
            f"Subject: {target}",
            f"Date: {self.report_date}",
            f"Risk Level: {risk.get('level', 'UNKNOWN')}",
            f"Risk Score: {risk.get('score', 0):.0f}/100",
            f"Investment Signal: {risk.get('invest_signal', 'UNKNOWN')}",
            "=" * 70,
        ]
        
        safe_name = target.replace(' ', '_')[:30]
        filepath = os.path.join(self.output_dir, f"Intel_Report_{safe_name}_{datetime.now().strftime('%Y%m%d')}.txt")
        
        with open(filepath, 'w') as f:
            f.write('\n'.join(lines))
        
        return filepath
